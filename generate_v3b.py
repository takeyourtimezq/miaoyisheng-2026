# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color

def header_bg(table, color_hex='4472C4'):
    for cell in table.rows[0].cells:
        s = OxmlElement('w:shd')
        s.set(qn('w:fill'), color_hex)
        s.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(s)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)

def borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    b = OxmlElement('w:tblBorders')
    for n in ['top','left','bottom','right','insideH','insideV']:
        e = OxmlElement(f'w:{n}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), '999999')
        b.append(e)
    tblPr.append(b)

def hl_row(row, color='E7F3FF'):
    for cell in row.cells:
        s = OxmlElement('w:shd')
        s.set(qn('w:fill'), color)
        s.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(s)

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '微软雅黑'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def bold_text(doc, label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.bold = True
    p.add_run(text)

BRAND = '轻刻身材管理中心'

# ===== 封面 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'{BRAND}·乌鲁木齐门店选址分析')
r.font.size = Pt(22)
r.font.bold = True
r.font.name = '微软雅黑'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('30-50㎡瘦身塑形店型')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
t = doc.add_table(rows=4, cols=2, style='Table Grid')
borders(t)
for i, (k, v) in enumerate([('数据来源', '美团、高德地图'), ('门店类型', '30-50㎡瘦身塑形'), ('分析日期', '2026年5月21日'), ('版本', 'V3.0')]):
    set_cell(t.rows[i].cells[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(t.rows[i].cells[1], v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), 'F2F2F2')
    s.set(qn('w:val'), 'clear')
    t.rows[i].cells[0]._tc.get_or_add_tcPr().append(s)

# ===== 一、竞品商户分布 =====
doc.add_page_break()
heading(doc, '一、竞品商户分布', 1)

p = doc.add_paragraph('通过美团、高德地图搜索"瘦身""纤体""减肥""美容美体"等关键词，整理乌鲁木齐在营瘦身塑形类商户如下：')

t2 = doc.add_table(rows=13, cols=4, style='Table Grid')
borders(t2)
for j, h in enumerate(['序号', '商户名称', '所在商圈', '数据来源']):
    set_cell(t2.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
header_bg(t2)

shops = [
    ('1', '良卡瘦身', '铁路局（新市区宁波街）', '高德地图'),
    ('2', '太百购物减肥体验中心', '铁路局（太百购物4层）', '高德地图'),
    ('3', '锦炫阁美容美体中心', '铁路局（新市区十四街）', '高德地图'),
    ('4', '佳丽减肥美容店', '新市区', '高德地图'),
    ('5', '创鑫美业女子减肥会所', '新市区', '高德地图'),
    ('6', '专业减肥（西八家户路）', '新市区', '高德地图'),
    ('7', '高丽名人纤体美容中心', '友好北路（天一大厦）', '美团/高德'),
    ('8', '美道家智能美容馆', '友好路（友好时尚购物城）', '美团'),
    ('9', '婕斯抗衰老减肥美容生活馆', '红山（天山区时代广场）', '高德地图'),
    ('10', '芬芬美容·减肥名店', '南湖（水磨沟区南湖东路）', '美团'),
    ('11', '媚丽一生经络美容瘦身馆', '乌鲁木齐（地址不详）', '高德地图'),
    ('12', '青橙瘦身塑形训练营', '乌鲁木齐（连锁品牌）', '美团'),
]
for i, row in enumerate(shops):
    for j, v in enumerate(row):
        set_cell(t2.rows[i+1].cells[j], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()

# 区域汇总
heading(doc, '商圈竞品密度汇总', 2)

t3 = doc.add_table(rows=6, cols=3, style='Table Grid')
borders(t3)
for j, h in enumerate(['商圈', '竞品数量', '市场判断']):
    set_cell(t3.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
header_bg(t3)

areas = [
    ('铁路局/新市区', '6家', '竞品最密集，市场需求已验证'),
    ('友好北路/友好路', '2家', '中高端定位，租金高'),
    ('红山/天山区', '1-2家', '客群偏年长，匹配度一般'),
    ('南湖/水磨沟区', '1-2家', '有市场但分散'),
    ('新疆民俗街/新华南路', '0家', '蓝海，无直接竞品'),
]
for i, row in enumerate(areas):
    for j, v in enumerate(row):
        set_cell(t3.rows[i+1].cells[j], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
hl_row(t3.rows[1])

# ===== 二、重点商圈分析 =====
doc.add_page_break()
heading(doc, '二、重点商圈分析（30-50㎡店型）', 1)

# 铁路局
heading(doc, '1. 铁路局商圈 ⭐ 首选', 2)
doc.add_paragraph('乌鲁木齐新市区老牌生活商圈，以铁路局职工社区为核心，辐射周边3-5公里居民区。核心商业体包括太百购物（YoungPark）等。')
bold_text(doc, '竞品情况：', '6家瘦身塑形类商户在营，为全城竞品最密集区域。太百购物内已有减肥体验中心入驻，说明商场内铺位可行。')
bold_text(doc, '客群画像：', '30-45岁中产家庭女性为主，消费稳定，健康意识强，有减脂/塑形需求。')
bold_text(doc, '租金水平：', '预估80-150元/㎡/月，30-50㎡月租金约2,400-7,500元。')
bold_text(doc, '优势：', '市场已验证、铺型匹配度高（30-80㎡为主流面积段）、成本低、社区口碑传播强。')
bold_text(doc, '风险：', '已有竞品先发优势，需差异化切入；商圈调性偏传统，品牌形象需提升。')

# 友好路
heading(doc, '2. 友好北路/友好路商圈', 2)
doc.add_paragraph('乌鲁木齐核心商业地标，聚集美美百货、友好商场、友好时尚购物城、汇嘉时代等高端商业体。')
bold_text(doc, '竞品情况：', '2家中高端纤体机构（高丽名人、美道家），定位偏高端，客单价较高。')
bold_text(doc, '租金水平：', '预估150-300元/㎡/月，30-50㎡月租金约4,500-15,000元。')
bold_text(doc, '判断：', '流量大、客群质量高，但租金和竞争门槛高，更适合作为品牌旗舰店，待样板店跑通后再入驻。')

# 民俗街
heading(doc, '3. 新疆民俗街/新华南路商圈', 2)
doc.add_paragraph('以新疆国际大巴扎为核心的文旅商圈，日均游客量大，夜间经济活跃。')
bold_text(doc, '竞品情况：', '瘦身塑形类商户为零，属于蓝海市场。')
bold_text(doc, '判断：', '有差异化机会（旅游+减脂体验），但游客转化率不确定、季节性波动大，建议作为第二梯队评估。')

# 其他
heading(doc, '4. 其他商圈', 2)

t4 = doc.add_table(rows=5, cols=3, style='Table Grid')
borders(t4)
for j, h in enumerate(['商圈', '竞品', '建议']):
    set_cell(t4.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
header_bg(t4)
for i, row in enumerate([
    ('北京南路', '1-2家', '备选社区店，优先级低于铁路局'),
    ('水上乐园/南湖', '1-2家', '季节性明显，不推荐首店'),
    ('红山/天山区', '1-2家', '客群年龄偏大，匹配度一般'),
    ('白鸟湖/乌北站/粮校', '0家', '人口密度低，商业不成熟，不推荐'),
]):
    for j, v in enumerate(row):
        set_cell(t4.rows[i+1].cells[j], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

# ===== 三、总结 =====
doc.add_page_break()
heading(doc, '三、总结', 1)

t7 = doc.add_table(rows=1, cols=1, style='Table Grid')
borders(t7)
s = OxmlElement('w:shd')
s.set(qn('w:fill'), 'FFF8E1')
s.set(qn('w:val'), 'clear')
t7.rows[0].cells[0]._tc.get_or_add_tcPr().append(s)

summary = (
    f'基于美团和高德地图竞品数据，乌鲁木齐瘦身塑形类商户共梳理出12家，其中铁路局/新市区占6家（50%），'
    '为全城竞品最密集区域。竞品集中不等于不能做——恰恰说明该区域市场需求真实、已被验证。'
    '\n\n'
    f'{BRAND} 30-50㎡瘦身塑形首店，建议选址铁路局商圈。'
    '核心原因：①市场已验证（6家竞品持续在营）；②成本可控（月租金2,400-7,500元）；'
    '③铺型匹配（30-50㎡为该商圈主流铺型）。'
    '\n\n'
    '友好路商圈留作品牌旗舰店（样板店跑通后），民俗街商圈作为差异化选项观察。'
)
set_cell(t7.rows[0].cells[0], summary, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT)

path = f'/Users/flyingspur/.qclaw/workspace-agent-6f07984c/{BRAND}_乌鲁木齐选址分析_V3.docx'
doc.save(path)
print(f'OK: {path}')
