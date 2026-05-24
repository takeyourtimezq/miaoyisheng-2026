#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 封面
title = doc.add_heading('升单爆破训练营', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.add_run('—— 美业连锁品牌升单特训 ——').bold = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

subtitle2 = doc.add_paragraph('招商方案')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle3 = doc.add_paragraph('线上培训 + 1个月陪跑')
subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle3.runs[0].font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n')

footer = doc.add_paragraph('适用对象：美业连锁品牌代理商/门店')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.color.rgb = RGBColor(153, 153, 153)

# 分页
doc.add_page_break()

# 目录
doc.add_heading('目录', 1)
toc = doc.add_paragraph()
toc.add_run('一、产品定位与价值\n'
'二、课程体系（2天培训）\n'
'三、陪跑机制（1个月SOP）\n'
'四、课程案例：39.9→99/200→1000-2000\n'
'五、价格策略\n'
'六、品牌方交付清单\n'
'七、话术手册摘要\n'
'八、我们的优势')

doc.add_paragraph('\n\n')

# 一、产品定位与价值
doc.add_heading('一、产品定位与价值', 1)
doc.add_heading('1.1 产品定位', 2)
p = doc.add_paragraph()
p.add_run('产品名称：').bold = True
p.add_run('升单爆破训练营')
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('目标客群：').bold = True
p2.add_run('美业连锁品牌代理商/门店店长/美容师')
doc.add_paragraph()
p3 = doc.add_paragraph()
p3.add_run('交付形式：').bold = True
p3.add_run('线上培训 + 1个月线上陪跑 / 线下培训 + 1个月线上陪跑')

doc.add_heading('1.2 核心价值', 2)
doc.add_paragraph('从"不会推"到"主动推"：解决升单心理障碍', style='List Number')
doc.add_paragraph('从"不敢推"到"轻松推"：掌握标准化话术武器库', style='List Number')
doc.add_paragraph('从"偶尔推"到"持续推"：1个月陪跑养成销售习惯', style='List Number')

# 二、课程体系
doc.add_heading('二、课程体系（2天培训）', 1)
doc.add_heading('2.1 课程架构', 2)
doc.add_paragraph('培训周期：2天（每天4-5小时）\n陪跑周期：1个月线上陪跑')

doc.add_heading('Day 1：开营 + 心法（3-4小时）', 3)
doc.add_paragraph('• 0-30min：开营仪式+训练营规则（30分钟）')
doc.add_paragraph('• 30-90min：升单底层逻辑：流量成本→必须升单（60分钟）')
doc.add_paragraph('• 90-150min：抖音客户消费心理：占便宜/试探/比较（60分钟）')
doc.add_paragraph('• 150-210min：接待SOP与破冰话术（60分钟）')
doc.add_paragraph('• 210-240min：小作业说明：竞品暗访（课后完成，30分钟）')

doc.add_heading('Day 2：话术 + 实战（4-5小时）', 3)
doc.add_paragraph('• 0-60min：需求挖掘：望闻问切+FABE话术（60分钟）')
doc.add_paragraph('• 60-120min：价格异议处理："太贵了"/"考虑一下"（60分钟）')
doc.add_paragraph('• 120-180min：升单时机：3个黄金节点（60分钟）')
doc.add_paragraph('• 180-240min：组合销售与套餐设计模板（60分钟）')
doc.add_paragraph('• 240-300min：场景演练通关：全流程模拟（60分钟）')
doc.add_paragraph('• 300-330min：陪跑启动+目标制定（30分钟）')

# 三、陪跑机制
doc.add_heading('三、陪跑机制（1个月SOP）', 1)
doc.add_heading('3.1 陪跑目标', 2)
doc.add_paragraph('巩固课程所学：将2天培训内容转化为日常销售习惯', style='List Number')
doc.add_paragraph('实战出结果：每学员每天至少3次升单尝试，目标升单率提升50%', style='List Number')
doc.add_paragraph('问题实时响应：每日答疑，快速解决实战中遇到的问题', style='List Number')

doc.add_heading('3.2 每日节奏', 2)
doc.add_paragraph('• 7:30 晨间分享：优秀案例/话术示范/激励语音（1分钟）')
doc.add_paragraph('• 12:00 今日任务下达：当日升单目标+学习任务')
doc.add_paragraph('• 18:00 数据上报：当天升单次数/成功/金额')
doc.add_paragraph('• 21:00 晚间复盘+答疑：当天问题汇总解答/优秀案例分享')
doc.add_paragraph('• 22:00 明日预告：明日课程/任务预告')

doc.add_heading('3.3 参与机制（消费者行为学设计）', 2)
doc.add_heading('解锁机制', 3)
doc.add_paragraph('当天任务完成→解锁次日课程回顾+新任务\n作业未通过→打回重做→解锁下一任务')

doc.add_heading('损失厌恶', 3)
doc.add_paragraph('连续3天未打卡→小组公示+班主任私聊\n连续7天未登录→移出陪跑群（可申请重入，需补交"违约金"）')

doc.add_heading('社交证明+荣誉体系', 3)
doc.add_paragraph('每日TOP3升单榜公示\n每周冠军小组奖励\n完成全部任务获得"升单高手"认证徽章')

doc.add_heading('小组PK制', 3)
doc.add_paragraph('5人一组，互相监督\n周冠军小组奖励：奖金/荣誉\n周垫底小组：表演节目/发红包')

doc.add_heading('3.4 周计划', 2)
doc.add_paragraph('第1周 破冰+需求挖掘：每天3次破冰尝试，记录客户反馈')
doc.add_paragraph('第2周 二转+异议处理：39.9→99/200二转尝试，价格异议处理练习')
doc.add_paragraph('第3周 利润品升单：99/200→1000+升单尝试，套餐组合练习')
doc.add_paragraph('第4周 综合实战+复盘：全流程实战，数据复盘，结业证书')

# 四、课程案例
doc.add_heading('四、课程案例：39.9→99/200→1000-2000', 1)
doc.add_heading('4.1 升单路径', 2)
p4 = doc.add_paragraph()
run = p4.add_run('抖音团购引流 39.9元 ')
run.bold = True
run.font.color.rgb = RGBColor(192, 0, 0)
doc.add_paragraph('↓')
p5 = doc.add_paragraph()
run2 = p5.add_run('到店体验 + 二转 99/200元 ')
run2.bold = True
run2.font.color.rgb = RGBColor(224, 96, 0)
doc.add_paragraph('↓')
p6 = doc.add_paragraph()
run3 = p6.add_run('利润品升单 1000-2000元 ')
run3.bold = True
run3.font.color.rgb = RGBColor(0, 176, 80)

doc.add_heading('4.2 全流程话术示例', 2)
doc.add_heading('第一步：39.9到店破冰（0-5分钟）', 3)
doc.add_paragraph('"您好！请坐～我是今天为您服务的美容师XX"\n"您是通过抖音过来的吗？第一次来我们店吧"\n"我们店主要做XX项目，您今天是想体验哪个呢？"')

doc.add_heading('第二步：需求挖掘（5-15分钟）', 3)
doc.add_paragraph('"您平时皮肤/头发有什么困扰吗？"\n"您希望改善到什么程度？"\n"我帮您分析一下，您这个问题是因为..."')

doc.add_heading('第三步：99元二转（15-30分钟）', 3)
doc.add_paragraph('"您今天体验的39.9项目是我们引流项目，非常划算"\n"我们店还有一个XX项目，性价比很高，99元3次"\n"这个项目正好能解决您刚才说的XX问题，我帮您预约一下？"')

doc.add_heading('第四步：1000-2000利润品升单（30-60分钟）', 3)
doc.add_paragraph('"您皮肤/头发的问题，需要搭配XX产品在家护理效果才持久"\n"我们店有XX套盒，效果很好，很多客户都在用"\n"今天正好有活动，买套盒送XX项目，全套1980元，比单买划算500元"')

# 五、价格策略
doc.add_heading('五、价格策略', 1)
doc.add_paragraph('注：以下价格为对品牌方收费（非对门店），包含品牌旗下所有门店学员')

# 价格表格
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '规模'
hdr_cells[1].text = '线上版'
hdr_cells[2].text = '线下版'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)

data = [
    ('30家门店', '9.8万', '19.8万'),
    ('50家门店', '15.8万', '32.8万'),
    ('100家门店', '28.8万', '58.8万'),
    ('200家+', '定制方案', '定制方案'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# 六、品牌方交付清单
doc.add_heading('六、品牌方交付清单', 1)

table2 = doc.add_table(rows=8, cols=3)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = '交付内容'
hdr_cells2[1].text = '线上版'
hdr_cells2[2].text = '线下版'
for cell in hdr_cells2:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)

data2 = [
    ('课程培训', '2天线上录播+直播', '1天线下集中授课'),
    ('线上陪跑', '1个月（30天）', '1个月（30天）'),
    ('学员账号', '30个', '30个'),
    ('督导名额', '1名线上督导', '1名线下督导+1名线上督导'),
    ('物料', '话术手册（电子）\n工具包（电子）', '话术手册（纸质+电子）\n工具包（纸质+电子）'),
    ('专属社群', '品牌专属陪跑群', '品牌专属陪跑群'),
    ('成果报告', '数据周报+结业报告', '门店诊断书+数据报告'),
]
for i, row_data in enumerate(data2):
    row = table2.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# 七、话术手册摘要
doc.add_heading('七、话术手册摘要', 1)
doc.add_heading('7.1 破冰话术（3句必杀）', 2)
doc.add_paragraph('1. "您好！请坐～我是今天为您服务的美容师XX"\n2. "您是通过抖音过来的吗？第一次来我们店吧"\n3. "我们店主要做XX项目，您今天是想体验哪个呢？"')

doc.add_heading('7.2 需求挖掘话术（望闻问切）', 2)
doc.add_paragraph('望：观察客户皮肤/头发状态\n闻：听客户描述问题\n问：您平时...？您希望...？\n切：给专业建议')

doc.add_heading('7.3 FABE话术模板', 2)
doc.add_paragraph('F（特性）：这个产品是...\nA（优势）：它可以...\nB（利益）：对您来说...\nE（证据）：很多客户用了都...（案例）')

doc.add_heading('7.4 价格异议处理（5种应对）', 2)
doc.add_paragraph('太贵了 → 分解法：算单次成本/对比线下/活动优惠', style='List Number')
doc.add_paragraph('考虑一下 → 紧迫感：活动限时/名额有限/下次涨价', style='List Number')
doc.add_paragraph('我考虑一下 → 追问法：您顾虑的是？', style='List Number')
doc.add_paragraph('别人更便宜 → 价值法：对比服务/产品/效果', style='List Number')
doc.add_paragraph('我要问老公 → 决策分离：先体验/先买先享受', style='List Number')

doc.add_heading('7.5 升单时机（3个黄金节点）', 2)
doc.add_paragraph('第一次项目做完（立即推）', style='List Number')
doc.add_paragraph('客户说"不错"（乘胜追击）', style='List Number')
doc.add_paragraph('客户问"多少钱"（组合销售）', style='List Number')

doc.add_heading('7.6 套餐设计模板', 2)
doc.add_paragraph('引流款：39.9元（抖音团单）→ 到店\n利润款：99元3次 → 绑定复购\n升级款：1000-2000元 → 盈利\n组合：买套盒送项目 = 价值感+绑定')

# 八、我们的优势
doc.add_heading('八、我们的优势', 1)
doc.add_heading('1. 专注美业，更懂美业', 2)
doc.add_paragraph('深耕美业连锁抖音本地生活，案例库丰富，实战经验丰富')

doc.add_heading('2. 理论+实战+陪跑三重保障', 2)
doc.add_paragraph('不仅讲理论，更注重实战；不仅培训，更陪跑到出结果')

doc.add_heading('3. 消费者行为学设计参与机制', 2)
doc.add_paragraph('课程解锁、损失厌恶、社交证明、荣誉体系，让学员主动学、持续练')

doc.add_heading('4. 可量化的结果承诺', 2)
doc.add_paragraph('陪跑结束提供数据报告，升单率提升效果可衡量')

# 联系
doc.add_paragraph('\n\n\n\n')
contact = doc.add_heading('联系我们', 1)
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('联系人：铖烈').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('电话：待补充').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
doc.save('/Users/flyingspur/.qclaw/workspace-agent-6f07984c/升单爆破训练营招商方案.docx')
print('Document created: 升单爆破训练营招商方案.docx')
