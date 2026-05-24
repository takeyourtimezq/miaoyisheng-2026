# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color

def header_bg(table, color_hex='4472C4'):
    for cell in table.rows[0].cells:
        s = OxmlElement('w:shd')
        s.set(qn('w:fill'), color_hex)
        s.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(s)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)

def borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    b = OxmlElement('w:tblBorders')
    for n in ['top','left','bottom','right','insideH','insideV']:
        e = OxmlElement(f'w:{n}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), '999999')
        b.append(e)
    tblPr.append(b)

def hl_row(row, color='E7F3FF'):
    for cell in row.cells:
        s = OxmlElement('w:shd')
        s.set(qn('w:fill'), color)
        s.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(s)

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '微软雅黑'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def bold_text(doc, label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.bold = True
    p.add_run(text)

BRAND = '轻刻身材管理中心'

# ===== 封面 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'{BRAND}·乌鲁木齐门店选址分析')
r.font.size = Pt(22)
r.font.bold = True
r.font.name = '微软雅黑'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('30-50㎡瘦身塑形店型')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
t = doc.add_table(rows=4, cols=2, style='Table Grid')
borders(t)
for i, (k, v) in enumerate([('数据来源', '抖音来客（到访热度）、美团、高德地图'), ('门店类型', '30-50㎡瘦身塑形'), ('分析日期', '2026年5月21日'), ('版本', 'V4.0')]):
    set_cell(t.rows[i].cells[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(t.rows[i].cells[1], v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), 'F2F2F2')
    s.set(qn('w:val'), 'clear')
    t.rows[i].cells[0]._tc.get_or_add_tcPr().append(s)

# ===== 一、抖音到访热度数据 =====
doc.add_page_break()
heading(doc, '一、抖音后台：预测到访商圈热度', 1)

p = doc.add_paragraph('抖音来客后台"预测TOP10到访商圈"，基于平台算法对目标客群实际到店可能性的预测排序：')

t1 = doc.add_table(rows=11, cols=3, style='Table Grid')
borders(t1)
for j, h in enumerate(['排名', '商圈', '到访热度']):
    set_cell(t1.rows[0].cells[j], h, bold=True, size=10, color=RGBColor(255,255,255))
header_bg(t1)

hot_data = [
    ('1', '铁路局', 'TOP1'),
    ('2', '新疆民俗街/新华南路', 'TOP2'),
    ('3', '水上乐园', 'TOP3'),
    ('4', '红山', 'TOP4'),
    ('5', '友好北路', 'TOP5'),
    ('6', '河南西路', 'TOP6'),
    ('7', '乌北站', 'TOP7'),
    ('8', '喀什西路', 'TOP8'),
    ('9', '粮校', 'TOP9'),
    ('10', '自治区博物馆', 'TOP10'),
]
for i, row in enumerate(hot_data):
    for j, v in enumerate(row):
        set_cell(t1.rows[i+1].cells[j], v, size=10)
hl_row(t1.rows[1])

# ===== 二、美团/高德竞品分布 =====
heading(doc, '二、美团/高德地图：竞品商户分布', 1)

p = doc.add_paragraph('通过美团、高德地图搜索"瘦身""纤体""减肥""美容美体"等关键词，乌鲁木齐在营瘦身塑形类商户如下：')

t2 = doc.add_table(rows=13, cols=4, style='Table Grid')
borders(t2)
for j, h in enumerate(['序号', '商户名称', '所在商圈', '数据来源']):
    set_cell(t2.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
header_bg(t2)

shops = [
    ('1', '良卡瘦身', '铁路局（新市区宁波街）', '高德地图'),
    ('2', '太百购物减肥体验中心', '铁路局（太百购物4层）', '高德地图'),
    ('3', '锦炫阁美容美体中心', '铁路局（新市区十四街）', '高德地图'),
    ('4', '佳丽减肥美容店', '新市区', '高德地图'),
    ('5', '创鑫美业女子减肥会所', '新市区', '高德地图'),
    ('6', '专业减肥（西八家户路）', '新市区', '高德地图'),
    ('7', '高丽名人纤体美容中心', '友好北路（天一大厦）', '美团/高德'),
    ('8', '美道家智能美容馆', '友好路（友好时尚购物城）', '美团'),
    ('9', '婕斯抗衰老减肥美容生活馆', '红山（天山区时代广场）', '高德地图'),
    ('10', '芬芬美容·减肥名店', '南湖（水磨沟区南湖东路）', '美团'),
    ('11', '媚丽一生经络美容瘦身馆', '乌鲁木齐（地址不详）', '高德地图'),
    ('12', '青橙瘦身塑形训练营', '乌鲁木齐（连锁品牌）', '美团'),
]
for i, row in enumerate(shops):
    for j, v in enumerate(row):
        set_cell(t2.rows[i+1].cells[j], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

# ===== 三、交叉验证分析 =====
doc.add_page_break()
heading(doc, '三、交叉验证分析', 1)

p = doc.add_paragraph('将抖音到访热度与美团/高德竞品分布叠加对比，得出各商圈的综合判断：')

t3 = doc.add_table(rows=8, cols=5, style='Table Grid')
borders(t3)
headers3 = ['商圈', '抖音到访热度', '竞品数量', '供需关系', '综合评级']
for j, h in enumerate(headers3):
    set_cell(t3.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
header_bg(t3)

cross_data = [
    ('铁路局', 'TOP1', '6家+', '需求大+供给足=市场成熟', '⭐⭐⭐⭐⭐ 首选'),
    ('新疆民俗街/新华南路', 'TOP2', '0家', '需求大+无供给=蓝海机会', '⭐⭐⭐⭐ 差异化选择'),
    ('友好北路', 'TOP5', '2家', '需求中等+高端供给=门槛高', '⭐⭐⭐ 备选旗舰店'),
    ('红山', 'TOP4', '1-2家', '需求中等+有供给=一般', '⭐⭐ 不推荐首店'),
    ('水上乐园', 'TOP3', '1-2家', '需求高但季节性强', '⭐⭐ 不推荐首店'),
    ('河南西路/喀什西路', 'TOP6/TOP8', '0-1家', '需求偏低+商业不成熟', '⭐ 不推荐'),
    ('乌北站/粮校/博物馆', 'TOP7-10', '0家', '需求低+边缘区域', '⭐ 不推荐'),
]
for i, row in enumerate(cross_data):
    for j, v in enumerate(row):
        set_cell(t3.rows[i+1].cells[j], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
hl_row(t3.rows[1])

doc.add_paragraph()

# 关键发现
bold_text(doc, '关键发现一：铁路局——三重验证', '')
doc.add_paragraph('抖音到访热度TOP1 + 美团/高德竞品最密集（6家+）+ 铁路局本身就是新市区核心生活区。三个独立信号指向同一个结论：这个区域的瘦身塑形市场需求真实存在且已被持续验证。竞品多不是坏事，说明有人愿意为这类服务买单。')

bold_text(doc, '关键发现二：民俗街——供需错配', '')
doc.add_paragraph('抖音到访热度TOP2，但美团/高德上几乎找不到同类商户。这说明该区域有明确的到访需求，但市场上没有满足需求的供给。属于蓝海机会，但需注意民俗街客流以游客为主，转化逻辑和社区型门店不同。')

bold_text(doc, '关键发现三：友好路——错位竞争', '')
doc.add_paragraph('抖音到访热度仅排TOP5（低于预期），但竞品定位偏高端（高丽名人纤体等）。说明友好路的流量更偏向高消费力人群，适合品牌势能建立后的旗舰店布局，不适合作为第一家样板店。')

# ===== 四、重点商圈详细分析 =====
doc.add_page_break()
heading(doc, '四、重点商圈分析（30-50㎡店型）', 1)

# 铁路局
heading(doc, '1. 铁路局商圈 ⭐ 首选', 2)
doc.add_paragraph('乌鲁木齐新市区老牌生活商圈，以铁路局职工社区为核心，辐射周边3-5公里居民区。核心商业体包括太百购物（YoungPark）等。')
bold_text(doc, '抖音数据：', '预测到访热度全城TOP1，目标客群最有可能实际到店的区域。')
bold_text(doc, '竞品情况：', '6家以上瘦身塑形类商户在营（良卡瘦身、太百减肥体验中心、锦炫阁美容美体、佳丽减肥美容、创鑫美业女子减肥会所、专业减肥），全城最密集。')
bold_text(doc, '客群画像：', '30-45岁中产家庭女性为主，消费稳定，健康意识强，有减脂/塑形需求。')
bold_text(doc, '租金水平：', '预估80-150元/㎡/月，30-50㎡月租金约2,400-7,500元。')
bold_text(doc, '优势：', '三重验证（抖音TOP1+竞品最密+成熟社区）、铺型匹配度高（30-80㎡为主流面积段）、成本低、社区口碑传播强。')
bold_text(doc, '风险：', '已有竞品先发优势，需差异化切入；商圈调性偏传统，品牌形象需提升。')

# 民俗街
heading(doc, '2. 新疆民俗街/新华南路商圈 ⭐ 差异化选择', 2)
doc.add_paragraph('以新疆国际大巴扎为核心的文旅商圈，日均游客量大，夜间经济活跃。')
bold_text(doc, '抖音数据：', '预测到访热度TOP2，仅次于铁路局。')
bold_text(doc, '竞品情况：', '瘦身塑形类商户为零，蓝海市场。')
bold_text(doc, '机会点：', '供需错配——有到访需求但没有服务供给。可考虑差异化打法（旅游打卡+减脂体验），利用大巴扎网红属性做内容引流。')
bold_text(doc, '风险：', '游客转化率不确定、季节性波动大、民族特色商圈品牌调性需适配。建议作为第二梯队评估。')

# 友好路
heading(doc, '3. 友好北路商圈 ⭐ 备选旗舰店', 2)
doc.add_paragraph('乌鲁木齐核心商业地标，聚集美美百货、友好商场、友好时尚购物城等高端商业体。')
bold_text(doc, '抖音数据：', '预测到访热度TOP5，低于铁路局和民俗街。')
bold_text(doc, '竞品情况：', '2家中高端纤体机构（高丽名人、美道家），定位偏高端。')
bold_text(doc, '判断：', '流量质量高但数量不是最高，租金和竞争门槛都偏高。更适合作为品牌旗舰店，待铁路局样板店跑通后再入驻，用核心商圈势能为招商赋能。')

# 其他
heading(doc, '4. 其他商圈简要评价', 2)

t4 = doc.add_table(rows=5, cols=3, style='Table Grid')
borders(t4)
for j, h in enumerate(['商圈', '抖音热度', '建议']):
    set_cell(t4.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
header_bg(t4)
for i, row in enumerate([
    ('水上乐园', 'TOP3', '季节性明显，不推荐首店'),
    ('红山', 'TOP4', '客群年龄偏大，匹配度一般'),
    ('河南西路/喀什西路', 'TOP6/TOP8', '商业不成熟，不推荐'),
    ('乌北站/粮校/博物馆', 'TOP7-10', '边缘区域，不推荐'),
]):
    for j, v in enumerate(row):
        set_cell(t4.rows[i+1].cells[j], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

# ===== 五、总结 =====
doc.add_page_break()
heading(doc, '五、总结', 1)

t7 = doc.add_table(rows=1, cols=1, style='Table Grid')
borders(t7)
s = OxmlElement('w:shd')
s.set(qn('w:fill'), 'FFF8E1')
s.set(qn('w:val'), 'clear')
t7.rows[0].cells[0]._tc.get_or_add_tcPr().append(s)

summary = (
    f'本次分析采用两组独立数据源交叉验证：①抖音来客"预测TOP10到访商圈"反映目标客群的实际到店意愿；'
    '②美团/高德地图竞品分布反映线下市场的真实供给格局。\n\n'
    f'{BRAND} 30-50㎡瘦身塑形首店，建议选址铁路局商圈。'
    '\n\n'
    '理由：铁路局在抖音到访热度中排名第一，同时是美团/高德上瘦身塑形竞品最密集的区域（6家+）。'
    '两个维度的数据形成互证——有真实的到访需求，也有持续经营的供给方，说明这是一个已经被市场验证过的成熟区域。'
    '30-50㎡在该商圈属于主流铺型，月租金可控（2,400-7,500元），试错成本合理。\n\n'
    '备选方案：新疆民俗街/新华南路（抖音到访TOP2但零竞品，蓝海机会）；'
    '友好北路（留作品牌旗舰店，待样板店跑通后入驻）。'
)
set_cell(t7.rows[0].cells[0], summary, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT)

path = f'/Users/flyingspur/.qclaw/workspace-agent-6f07984c/{BRAND}_乌鲁木齐选址分析_V4.docx'
doc.save(path)
print(f'OK: {path}')
