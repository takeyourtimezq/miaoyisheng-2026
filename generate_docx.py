# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell_font(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color

def set_header_row(table, color_hex='4472C4'):
    for cell in table.rows[0].cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '999999')
        borders.append(border)
    tblPr.append(borders)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

# ========== 封面信息 ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('轻刻时代·乌鲁木齐门店选址分析报告')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— 30-50㎡瘦身塑形店型专项分析 ——')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# 基本信息表
info = doc.add_table(rows=5, cols=2, style='Table Grid')
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('分析项目', '轻刻时代乌鲁木齐首店选址'),
    ('门店类型', '30-50㎡瘦身塑形门店'),
    ('数据来源', '抖音来客 + 美团/大众点评 + 高德地图 + 天眼查'),
    ('分析日期', '2026年5月21日'),
    ('报告版本', 'V2.0 交叉验证版'),
]
for i, (k, v) in enumerate(info_data):
    set_cell_font(info.rows[i].cells[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_font(info.rows[i].cells[1], v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    # 灰色背景
    for cell in [info.rows[i].cells[0]]:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F2F2F2')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

doc.add_page_break()

# ========== 一、分析背景与方法 ==========
add_heading_styled(doc, '一、分析背景与方法', level=1)

doc.add_paragraph('轻刻时代作为健康减脂品牌，计划在乌鲁木齐开设首家30-50㎡瘦身塑形门店。为确保选址决策有据可依，本次分析采用多数据源交叉验证方法：')

items = [
    '抖音来客数据 —— 行业兴趣人群地点偏好分布（直接反映目标客群的线上行为偏好）',
    '美团/大众点评 —— 搜索"瘦身""纤体""减肥""塑形"等关键词，提取已注册商户的地理位置分布（反映线下竞争格局）',
    '高德地图/360地图 —— 商铺搜索结果（反映商圈内美容美体类商户的聚集程度）',
    '天眼查/58同城 —— 企业注册信息与招聘信息（反映行业活跃度和经营状态）',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10.5)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('核心目标：')
run.font.bold = True
doc.add_paragraph('锁定1-2个最适合轻刻时代开30-50㎡瘦身塑形店的商圈，给出明确的选址优先级和建议动作。')

# ========== 二、数据交叉验证 ==========
doc.add_page_break()
add_heading_styled(doc, '二、数据交叉验证', level=1)

add_heading_styled(doc, '2.1 抖音来客：兴趣人群分布', level=2)

table1 = doc.add_table(rows=11, cols=4, style='Table Grid')
add_table_borders(table1)
headers = ['排名', '商圈名称', '兴趣人数', '数据占比']
for j, h in enumerate(headers):
    set_cell_font(table1.rows[0].cells[j], h, bold=True, size=10, color=RGBColor(255,255,255))
set_header_row(table1)

data_rows = [
    ('1', '铁路局', '8', '14.5%'),
    ('2', '新疆民俗街/新华南路', '6', '10.9%'),
    ('3', '友好北路', '5', '9.1%'),
    ('4', '水上乐园', '5', '9.1%'),
    ('5', '友好路', '5', '9.1%'),
    ('6', '北京南路', '5', '9.1%'),
    ('7', '红山', '4', '7.3%'),
    ('8', '白鸟湖', '4', '7.3%'),
    ('9', '乌北站', '4', '7.3%'),
    ('10', '粮校', '4', '7.3%'),
]
for i, row_data in enumerate(data_rows):
    for j, val in enumerate(row_data):
        set_cell_font(table1.rows[i+1].cells[j], val, size=10)
# 高亮TOP1
for cell in table1.rows[1].cells:
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E7F3FF')
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('说明：')
run.font.bold = True
p.add_run('合计55人样本。铁路局商圈以14.5%占比位列第一，说明该区域线上减脂兴趣人群最为集中。')

# 2.2 美团/大众点评竞品分布
doc.add_paragraph()
add_heading_styled(doc, '2.2 美团/大众点评 + 高德地图：竞品商户分布', level=2)

p = doc.add_paragraph()
run = p.add_run('通过美团、大众点评、高德地图、360地图、天眼查等平台搜索"瘦身""纤体""减肥""美容美体"关键词，整理乌鲁木齐已注册/在线经营的瘦身塑形类商户如下：')

table2 = doc.add_table(rows=13, cols=4, style='Table Grid')
add_table_borders(table2)
headers2 = ['序号', '商户名称', '所在区域', '数据来源']
for j, h in enumerate(headers2):
    set_cell_font(table2.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
set_header_row(table2)

competitors = [
    ('1', '良卡瘦身', '新市区（铁路局片区）宁波街', '联途/高德'),
    ('2', '铁路局减肥体验中心', '太百购物铁路局店4层', '微信公众号'),
    ('3', '锦炫阁美容美体中心', '新市区铁路局十四街', '天眼查'),
    ('4', '佳丽减肥美容店', '新市区', '天眼查'),
    ('5', '创鑫美业女子减肥会所', '新市区', '天眼查'),
    ('6', '高丽名人纤体美容中心', '沙依巴克区友好北路天一大厦', '天眼查/搜狗'),
    ('7', '美道家智能美容馆', '乌市区/昌吉区', '美团/微信公众号'),
    ('8', '婕斯抗衰老减肥美容生活馆', '天山区红山东路时代广场', '360地图'),
    ('9', '芬芬美容·减肥名店', '水磨沟区南湖东路', '360地图'),
    ('10', '媚丽一生经络美容瘦身馆', '乌鲁木齐市（地址不详）', '360地图'),
    ('11', '青橙瘦身塑形训练营', '乌鲁木齐（连锁品牌）', '抖音'),
    ('12', '专业减肥', '新市区西八家户路', '联途/高德'),
]
for i, row_data in enumerate(competitors):
    for j, val in enumerate(row_data):
        set_cell_font(table2.rows[i+1].cells[j], val, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()

# 竞品区域分布汇总
add_heading_styled(doc, '2.3 竞品区域分布汇总（交叉验证结论）', level=2)

table3 = doc.add_table(rows=6, cols=5, style='Table Grid')
add_table_borders(table3)
headers3 = ['商圈/区域', '抖音排名', '竞品数量', '竞争密度', '综合评价']
for j, h in enumerate(headers3):
    set_cell_font(table3.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
set_header_row(table3)

summary_data = [
    ('铁路局/新市区', 'TOP1（8人）', '5家+', '中等', '需求大+有竞品=市场已验证'),
    ('友好北路/友好路', 'TOP3+5（10人）', '2-3家', '较高', '流量大但租金高竞争强'),
    ('新疆民俗街/新华南路', 'TOP2（6人）', '0-1家', '低', '蓝海市场差异化机会'),
    ('红山/天山区', 'TOP7（4人）', '2-3家', '中等', '客群年龄偏大匹配度一般'),
    ('水上乐园/南湖', 'TOP4（5人）', '1-2家', '低', '季节性明显不推荐首店'),
]
for i, row_data in enumerate(summary_data):
    for j, val in enumerate(row_data):
        set_cell_font(table3.rows[i+1].cells[j], val, size=9)
# 高亮首选
for cell in table3.rows[1].cells:
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E7F3FF')
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('关键发现：')
run.font.bold = True
run.font.color.rgb = RGBColor(200, 50, 50)

findings = [
    '铁路局商圈在抖音数据（TOP1）和竞品分布（5家+，全城最多）上形成"双重验证"——说明该区域不是"有需求没市场"，而是"有需求、有市场、已验证"',
    '友好北路/友好路虽然流量最大（合并10人），但竞品为高端纤体机构（如高丽名人），租金和竞争门槛都高',
    '新疆民俗街/新华南路有抖音需求（6人）但竞品几乎为零，属于蓝海市场',
    '新市区是减肥瘦身类商户注册最密集的区域，侧面印证铁路局片区的市场成熟度',
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

# ========== 三、商圈详细分析 ==========
doc.add_page_break()
add_heading_styled(doc, '三、重点商圈详细分析（30-50㎡店型适配）', level=1)

# 铁路局商圈
add_heading_styled(doc, '3.1 铁路局商圈 ⭐ 首选', level=2)

p = doc.add_paragraph()
run = p.add_run('区域概况：')
run.font.bold = True
doc.add_paragraph('铁路局商圈位于乌鲁木齐新市区，是乌鲁木齐老牌生活商圈。以铁路局职工社区为核心，辐射周边3-5公里居民区。核心商业体包括太百购物铁路局店（YoungPark）、新市区各大社区商业街等。')

p = doc.add_paragraph()
run = p.add_run('交叉验证数据：')
run.font.bold = True
bullet_items = [
    '抖音兴趣人群：8人，占比14.5%，全城TOP1',
    '竞品商户：5家以上（良卡瘦身、太百减肥体验中心、锦炫阁美容美体、佳丽减肥美容、创鑫美业女子减肥会所等），为全城瘦身塑形商户最密集区域',
    '高德地图：搜索"瘦身""减肥"关键词，铁路局片区商户数量居首',
    '58同城：铁路局有美容美体类商铺出租信息，说明商业地产流通活跃',
    '天眼查：新市区减肥美容类个体工商户注册数量最多',
]
for item in bullet_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('30-50㎡店型适配分析：')
run.font.bold = True

table_fit = doc.add_table(rows=5, cols=2, style='Table Grid')
add_table_borders(table_fit)
fit_data = [
    ('评估维度', '分析'),
    ('面积适配', '30-50㎡属于小型门店，铁路局社区商铺常见面积段为30-80㎡，铺型匹配度高，可选面广'),
    ('租金成本', '预估80-150元/㎡/月，月租金约2,400-7,500元，在轻刻时代预算可控范围内'),
    ('竞品面积参考', '太百购物铁路局店4层已有减肥体验中心，说明商场内也有小型店面选项；街边店以30-60㎡为主'),
    ('客流量级', '社区日均人流量稳定，以周边3-5公里居民为主，适合"低频高客单"的瘦身塑形服务'),
]
for i, (k, v) in enumerate(fit_data):
    set_cell_font(table_fit.rows[i].cells[0], k, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(table_fit.rows[i].cells[1], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
    if i == 0:
        for cell in table_fit.rows[0].cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '4472C4')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('优势：')
run.font.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)
pros = [
    '需求已验证：5家以上竞品持续经营，说明市场需求真实存在',
    '铺型匹配：30-50㎡属于铁路局商圈主流铺型，选择空间大',
    '成本可控：月租金2,400-7,500元，试错成本低',
    '客群精准：中产家庭为主，30-45岁女性群体，减脂需求强',
    '口碑传播：老社区人际网络密集，利于转介绍',
    '线上引流配合：抖音兴趣人群TOP1，线上广告可以精准覆盖周边',
]
for item in pros:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('风险：')
run.font.bold = True
run.font.color.rgb = RGBColor(200, 50, 50)
cons = [
    '竞品存在：良卡瘦身等品牌已有先发优势，需差异化切入',
    '商圈调性偏传统：需在品牌形象和服务体验上做出区隔',
    '年轻白领较少：需通过抖音线上引流补充客群',
]
for item in cons:
    doc.add_paragraph(item, style='List Bullet')

# 友好路商圈
doc.add_page_break()
add_heading_styled(doc, '3.2 友好北路/友好路商圈 ⭐ 备选', level=2)

p = doc.add_paragraph()
run = p.add_run('区域概况：')
run.font.bold = True
doc.add_paragraph('友好北路/友好路是乌鲁木齐核心商业地标，聚集美美百货、友好商场、友好时尚购物城、汇嘉时代等高端商业体。是乌鲁木齐消费力最强、人流量最大的商圈。')

p = doc.add_paragraph()
run = p.add_run('交叉验证数据：')
run.font.bold = True
bullet_items = [
    '抖音兴趣人群：友好北路5人 + 友好路5人 = 10人，合并后全城TOP1',
    '竞品商户：高丽名人纤体美容中心（友好北路天一大厦）、美道家智能美容馆（友好时尚购物城负一楼）等',
    '竞品定位偏高端：高丽名人纤体定位高端、美道家定位智能美容，客单价偏高',
    '招聘信息：友好时尚购物城曾有美容机构招聘（月薪5000-7000+社保），说明该区域门店运营成本较高',
]
for item in bullet_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('30-50㎡店型适配分析：')
run.font.bold = True

doc.add_paragraph('友好路商圈以购物中心内铺位为主，独立街边30-50㎡铺位较少，且租金预估150-300元/㎡/月，月租金约4,500-15,000元。虽然流量大、客群质量高，但对于首店样板店来说成本偏高。')

p = doc.add_paragraph()
run = p.add_run('建议定位：')
run.font.bold = True
doc.add_paragraph('品牌旗舰店/招商展示店，待铁路局样板店跑通后再入驻，利用核心商圈势能为招商赋能。')

# 新疆民俗街商圈
add_heading_styled(doc, '3.3 新疆民俗街/新华南路商圈 ⭐ 差异化选择', level=2)

p = doc.add_paragraph()
run = p.add_run('区域概况：')
run.font.bold = True
doc.add_paragraph('以新疆国际大巴扎为核心，是乌鲁木齐最具代表性的文旅商圈。日均游客量大，夜间经济活跃，特色餐饮和文创商铺聚集。')

p = doc.add_paragraph()
run = p.add_run('交叉验证数据：')
run.font.bold = True
bullet_items = [
    '抖音兴趣人群：6人，占比10.9%，TOP2',
    '竞品商户：经美团/大众点评/高德地图搜索，该区域瘦身塑形类商户几乎为零',
    '属于蓝海市场：有需求（抖音验证）但无竞品，差异化机会明显',
]
for item in bullet_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('差异化打法建议：')
run.font.bold = True
doc.add_paragraph('主打"旅游打卡+减脂体验"双概念。大巴扎自带网红属性，适合做抖音内容创作引流。30-50㎡小型门店投入低，适合做轻资产试点。')

p = doc.add_paragraph()
run = p.add_run('风险：')
run.font.bold = True
doc.add_paragraph('游客转化率不确定性高、季节性波动大、民族特色商圈品牌调性需适配。建议作为第二梯队评估，不作为首店首选。')

# 其他商圈
add_heading_styled(doc, '3.4 其他商圈简要评价', level=2)

table4 = doc.add_table(rows=7, cols=4, style='Table Grid')
add_table_borders(table4)
headers4 = ['商圈', '抖音排名', '竞品情况', '30-50㎡店型建议']
for j, h in enumerate(headers4):
    set_cell_font(table4.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
set_header_row(table4)

other_data = [
    ('北京南路', 'TOP6（5人）', '1-2家', '备选社区店，优先级低于铁路局'),
    ('水上乐园/南湖', 'TOP4（5人）', '1-2家', '季节性明显，不推荐首店'),
    ('红山/天山区', 'TOP7（4人）', '2-3家', '客群年龄偏大，匹配度一般'),
    ('白鸟湖', 'TOP8（4人）', '0家', '新兴商圈，人口密度低，不推荐'),
    ('乌北站', 'TOP9（4人）', '0家', '边缘商圈，商业不成熟，不推荐'),
    ('粮校', 'TOP10（4人）', '0家', '边缘商圈，发展周期长，不推荐'),
]
for i, row_data in enumerate(other_data):
    for j, val in enumerate(row_data):
        set_cell_font(table4.rows[i+1].cells[j], val, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

# ========== 四、综合推荐 ==========
doc.add_page_break()
add_heading_styled(doc, '四、综合推荐与行动建议', level=1)

add_heading_styled(doc, '4.1 选址优先级排序', level=2)

table5 = doc.add_table(rows=4, cols=5, style='Table Grid')
add_table_borders(table5)
headers5 = ['优先级', '商圈', '推荐理由', '月租金预估', '适合阶段']
for j, h in enumerate(headers5):
    set_cell_font(table5.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
set_header_row(table5)

rec_data = [
    ('第一优先', '铁路局商圈', '抖音TOP1+竞品最密集=市场已验证', '2,400-7,500元', '0-1样板店'),
    ('第二优先', '友好北路/友好路', '流量最大+客群质量高', '4,500-15,000元', '品牌旗舰店'),
    ('第三优先', '新疆民俗街/新华南路', '蓝海市场+网红属性', '3,000-9,000元', '差异化试点'),
]
for i, row_data in enumerate(rec_data):
    for j, val in enumerate(row_data):
        set_cell_font(table5.rows[i+1].cells[j], val, size=9)
    # 高亮首选行
    if i == 0:
        for cell in table5.rows[1].cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'E7F3FF')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)

doc.add_paragraph()

add_heading_styled(doc, '4.2 铁路局商圈选铺建议', level=2)

doc.add_paragraph('基于30-50㎡店型需求，建议重点关注以下铺位类型：')

p = doc.add_paragraph()
run = p.add_run('街边底商（推荐指数：⭐⭐⭐⭐⭐）')
run.font.bold = True
doc.add_paragraph('首选铁路局商圈主干道（如南二路、北京南路等）临街商铺。面积30-50㎡，门头宽度≥3米，层高≥2.8米。需满足以下条件：')
conditions = [
    '周边500米内有3个以上成熟小区',
    '可视性好，有独立招牌位',
    '上水下水到位，可接三相电',
    '有停车位或周边停车便利',
]
for c in conditions:
    doc.add_paragraph(c, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('社区商业体/购物中心内铺（推荐指数：⭐⭐⭐⭐）')
run.font.bold = True
doc.add_paragraph('太百购物铁路局店（YoungPark）已有减肥体验中心入驻先例，可考虑同类型商业体内的30-50㎡内铺。优势：自带客流、物业管理规范、品牌形象好。')

p = doc.add_paragraph()
run = p.add_run('写字楼底商（推荐指数：⭐⭐⭐）')
run.font.bold = True
doc.add_paragraph('铁路局片区写字楼底商租金较低，适合预算紧张时考虑。但需注意：白领客群偏年轻，需配合抖音引流补充社区客群。')

add_heading_styled(doc, '4.3 启动成本粗算', level=2)

table6 = doc.add_table(rows=8, cols=2, style='Table Grid')
add_table_borders(table6)
set_cell_font(table6.rows[0].cells[0], '成本项目', bold=True, size=9, color=RGBColor(255,255,255))
set_cell_font(table6.rows[0].cells[1], '预估金额', bold=True, size=9, color=RGBColor(255,255,255))
set_header_row(table6)

cost_data = [
    ('押金（押二付一或押三付一）', '5,000-22,500元'),
    ('装修（30-50㎡，轻装修）', '30,000-80,000元'),
    ('设备（瘦身仪器+基础家具）', '20,000-50,000元'),
    ('证照办理', '2,000-5,000元'),
    ('首月人员工资（1-2人）', '5,000-14,000元'),
    ('首批物料/耗材', '3,000-8,000元'),
    ('合计启动资金', '65,000-179,500元'),
]
for i, (k, v) in enumerate(cost_data):
    set_cell_font(table6.rows[i+1].cells[0], k, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(table6.rows[i+1].cells[1], v, size=9)
    if i == len(cost_data) - 1:
        for cell in table6.rows[i+1].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True

add_heading_styled(doc, '4.4 下一步行动清单', level=2)

table7 = doc.add_table(rows=7, cols=3, style='Table Grid')
add_table_borders(table7)
headers7 = ['序号', '行动事项', '时间建议']
for j, h in enumerate(headers7):
    set_cell_font(table7.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
set_header_row(table7)

actions = [
    ('1', '实地考察铁路局商圈，重点看南二路、北京南路、太百购物周边铺位', '本周内'),
    ('2', '走访铁路局商圈竞品（良卡瘦身、太百减肥体验中心），了解其客单价、客流量、服务项目', '考察期间'),
    ('3', '通过58同城/安居客筛选铁路局30-50㎡商铺出租信息', '本周内'),
    ('4', '基于2-3个候选铺位，测算详细启动成本和月运营成本', '考察后3天'),
    ('5', '预估盈亏平衡点（假设客单价、月到店人数、复购率）', '测算后2天'),
    ('6', '了解乌鲁木齐健康减脂/美容类经营许可证办理流程', '签约前'),
]
for i, row_data in enumerate(actions):
    for j, val in enumerate(row_data):
        set_cell_font(table7.rows[i+1].cells[j], val, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

# ========== 五、免责声明 ==========
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('免责声明')
run.font.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)

p = doc.add_paragraph()
run = p.add_run('本报告基于抖音来客后台数据、美团/大众点评公开信息、高德地图/360地图搜索结果、天眼查企业信息等公开数据源交叉验证分析。数据采集时间为2026年5月21日，商户信息可能存在时效性偏差。租金预估基于乌鲁木齐同类商圈公开数据推算，实际租金需以实地调研和房东报价为准。本报告仅供参考，实际选址决策需结合实地调研及专业评估。')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(150, 150, 150)

# 保存
output_path = '/Users/flyingspur/.qclaw/workspace-agent-6f07984c/轻刻时代_乌鲁木齐选址分析报告_V2.docx'
doc.save(output_path)
print(f'Word文档生成成功：{output_path}')
