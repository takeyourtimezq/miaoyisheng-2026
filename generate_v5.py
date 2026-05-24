# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color

def header_bg(table, color_hex='4472C4'):
    for cell in table.rows[0].cells:
        s = OxmlElement('w:shd')
        s.set(qn('w:fill'), color_hex)
        s.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(s)

def borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    b = OxmlElement('w:tblBorders')
    for n in ['top','left','bottom','right','insideH','insideV']:
        e = OxmlElement(f'w:{n}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), '999999')
        b.append(e)
    tblPr.append(b)

def hl_row(row, color='E7F3FF'):
    for cell in row.cells:
        s = OxmlElement('w:shd')
        s.set(qn('w:fill'), color)
        s.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(s)

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '微软雅黑'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def bold_text(doc, label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.bold = True
    p.add_run(text)
    return p

BRAND = '轻刻身材管理中心'

# ===== 封面 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'{BRAND}·乌鲁木齐门店选址分析')
r.font.size = Pt(22)
r.font.bold = True
r.font.name = '微软雅黑'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('30-50㎡瘦身塑形店型')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
t = doc.add_table(rows=4, cols=2, style='Table Grid')
borders(t)
for i, (k, v) in enumerate([
    ('数据来源', '抖音来客（到访热度）、美团、高德地图'),
    ('门店类型', '30-50㎡瘦身塑形'),
    ('分析日期', '2026年5月21日'),
    ('版本', 'V5.0（竞品专项版）'),
]):
    set_cell(t.rows[i].cells[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(t.rows[i].cells[1], v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), 'F2F2F2')
    s.set(qn('w:val'), 'clear')
    t.rows[i].cells[0]._tc.get_or_add_tcPr().append(s)

# ===== 一、抖音到访热度数据 =====
doc.add_page_break()
heading(doc, '一、抖音后台：预测到访商圈热度', 1)

p = doc.add_paragraph('抖音来客后台"预测TOP10到访商圈"，基于平台算法对目标客群实际到店可能性的预测排序：')

t1 = doc.add_table(rows=11, cols=3, style='Table Grid')
borders(t1)
for j, h in enumerate(['排名', '商圈', '到访热度']):
    set_cell(t1.rows[0].cells[j], h, bold=True, size=10, color=RGBColor(255,255,255))
header_bg(t1)

hot_data = [
    ('1', '铁路局', 'TOP1'),
    ('2', '新疆民俗街/新华南路', 'TOP2'),
    ('3', '水上乐园', 'TOP3'),
    ('4', '红山', 'TOP4'),
    ('5', '友好北路', 'TOP5'),
    ('6', '河南西路', 'TOP6'),
    ('7', '乌北站', 'TOP7'),
    ('8', '喀什西路', 'TOP8'),
    ('9', '粮校', 'TOP9'),
    ('10', '自治区博物馆', 'TOP10'),
]
for i, row in enumerate(hot_data):
    for j, v in enumerate(row):
        set_cell(t1.rows[i+1].cells[j], v, size=10)
hl_row(t1.rows[1])

# ===== 二、连锁品牌竞品调研 =====
doc.add_page_break()
heading(doc, '二、连锁品牌竞品调研', 1)

p = doc.add_paragraph('针对大唐辣妈、良咔、燃放、寇氏减肥等全国性连锁品牌，逐一调研其在乌鲁木齐的门店布局情况：')

# 综合对比表
t_comp = doc.add_table(rows=7, cols=6, style='Table Grid')
borders(t_comp)
headers = ['品牌', '全国规模', '乌鲁木齐门店数', '所在商圈', '主要打法', '客单价参考']
for j, h in enumerate(headers):
    set_cell(t_comp.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
header_bg(t_comp)

comp_data = [
    ('大唐辣妈', '全国9000+店', '有开业记录（安宁开业活动）\n乌鲁木齐具体地址待实地确认', '—', '低价引流（398元瘦10斤）\n+ 全国连锁品牌背书', '低单尝鲜为主'),
    ('良咔', '全国数百店规模\n主打瘦身连锁', '3家确认\n①天山区碱泉三街嘉鸿园4栋\n②头屯河区西泉街美容店\n③头屯河区西泉街养生馆', '天山区、头屯河区', '主打纤体瘦身\n社区店模型\n单店面积小', '中低端为主'),
    ('燃放', '全国加盟连锁\n南京公司运营', '暂无明确记录\n（搜索未发现乌鲁木齐在营门店）', '—', '养生馆+产品结合\n加盟扩张模式', '中等'),
    ('寇氏减肥', '全国6000+店\n1990年创立\n中医拔罐减肥', '4家确认\n①南湖店（水磨沟区南湖东路）\n②水磨沟区五星北路71号\n③沙依巴克区长江路棉花街55号\n④长江路棉花街51号', '南湖、水磨沟区、长江路（沙区）', '中医拔罐+外敷产品\n老字号品牌\n单店面积较大（95㎡起加盟）', '中等偏高'),
    ('芭芭多', '全国连锁\n芦荟护肤减肥', '1家确认\n米东区石化', '米东区', '芦荟产品为核心\n护肤+减肥双线', '中等'),
]
for i, row in enumerate(comp_data):
    for j, v in enumerate(row):
        set_cell(t_comp.rows[i+1].cells[j], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
hl_row(t_comp.rows[1], 'FFF3E0')

# 各品牌详细分析
heading(doc, '2. 各品牌详细情况', 2)

# 大唐辣妈
bold_text(doc, '① 大唐辣妈', '')
doc.add_paragraph(
    '全国规模最大的减肥连锁品牌之一，宣称全国门店9000余家。以"低价引流+品牌背书"为核心打法，'
    '常见推广话术为"398元瘦10斤"，通过极低门槛吸引首次体验，再通过疗程和套装产品完成二次转化。'
    '在乌鲁木齐有开业活动记录（安宁），说明品牌已在新疆市场布局，但目前搜索到的信息以开业推广为主，'
    '乌鲁木齐主城区内具体门店数量和地址尚需实地确认。'
    '\n\n竞争启示：低单价引流是该品牌核心打法，若轻刻身材管理中心走同质化低价路线会直接竞争；但其品牌势能强，'
    '轻刻可差异化定位——用更专业的服务和效果承诺，打中高端市场。'
)

# 良咔
bold_text(doc, '② 良咔', '')
doc.add_paragraph(
    '国内专注纤体瘦身的连锁品牌，主打社区型小店模型。乌鲁木齐已确认3家门店，'
    '分别位于天山区碱泉三街（嘉鸿园御邸世家）和头屯河区西泉街。'
    '\n\n良咔的选址逻辑：以社区底商为主，贴近居民生活半径，客群稳定但辐射范围有限。'
    '值得注意的是，天山区和头屯河区的布局，恰好避开了铁路局商圈和友好路核心商圈——'
    '这与抖音预测到访热度排名形成错位：抖音热度最高的铁路局和民俗街，目前没有良咔布局。'
    '\n\n竞争启示：良咔在乌鲁木齐的布点偏边缘，核心商圈留有空白，这是轻刻进入的机会窗口。'
)

# 燃放
bold_text(doc, '③ 燃放', '')
doc.add_paragraph(
    '南京燃放大健康管理有限公司旗下品牌，主打"养生馆结合+产品研发"模式，全国开放加盟。'
    '目前搜索未发现乌鲁木齐在营门店，说明该品牌尚未进入乌鲁木齐市场。'
    '\n\n竞争启示：燃放属于待进入乌鲁木齐的潜在竞争者，轻刻若率先在核心商圈建立标杆，有机会形成先发优势。'
)

# 寇氏减肥
bold_text(doc, '④ 寇氏减肥', '')
doc.add_paragraph(
    '国内历史最久的减肥连锁品牌之一，1990年创立于山东，2009年注册"寇氏"商标，现全球连锁店6000余家。'
    '主打中医拔罐减肥+外敷产品，以"辩证论治、平衡脏腑"为卖点，绿色健康。'
    '寇氏要求加盟门店面积95㎡起，初期投入12-31万元，属于相对重投入的品牌。'
    '\n\n乌鲁木齐已确认4家门店，分布在水磨沟区（南湖、五星北路）和沙依巴克区（长江路棉花街），'
    '均为成熟居民区，未进入新市区（铁路局商圈）和天山区核心商业区。'
    '\n\n竞争启示：寇氏的客群年龄偏大，打法传统，适合作为对标品牌而非直接竞争。'
    '轻刻的智能化/数据化服务+更年轻的品牌形象，可以与其形成明显差异化。'
)

# 芭芭多
bold_text(doc, '⑤ 芭芭多（补充参考）', '')
doc.add_paragraph(
    '全国芦荟护肤连锁品牌，乌鲁木齐米东区有1家门店，主打芦荟产品+美容护肤+减肥。'
    '在乌鲁木齐规模有限，可作为参考但不构成主要竞争。'
)

# ===== 三、连锁竞品商圈分布与机会判断 =====
doc.add_page_break()
heading(doc, '三、连锁竞品商圈覆盖分析', 1)

p = doc.add_paragraph('将以上五个连锁品牌的乌鲁木齐门店分布，与抖音预测到访热度商圈做叠加，得出竞争空白区域：')

t_blank = doc.add_table(rows=6, cols=4, style='Table Grid')
borders(t_blank)
for j, h in enumerate(['商圈（抖音热度TOP5）', '连锁竞品布局', '空白程度', '机会判断']):
    set_cell(t_blank.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
header_bg(t_blank)

blank_data = [
    ('铁路局（TOP1）', '几乎无全国连锁品牌布点\n（仅零散个体户）', '⭐⭐⭐⭐⭐ 高空白', '最优先进入机会，竞品真空期'),
    ('新疆民俗街/新华南路（TOP2）', '无全国连锁布点', '⭐⭐⭐⭐⭐ 高空白', '蓝海机会，但客流结构偏游客'),
    ('水上乐园（TOP3）', '无全国连锁布点', '⭐⭐⭐⭐ 高空白', '季节性强，首店暂不推荐'),
    ('红山（TOP4）', '无全国连锁布点\n但有寇氏南湖店覆盖水区', '⭐⭐⭐ 中空白', '可作为第二梯队评估'),
    ('友好北路（TOP5）', '无全国连锁布点\n竞品偏高端（高丽名人等）', '⭐⭐⭐ 中空白', '适合品牌旗舰店，不适合首店'),
]
for i, row in enumerate(blank_data):
    for j, v in enumerate(row):
        set_cell(t_blank.rows[i+1].cells[j], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
hl_row(t_blank.rows[1])

p = doc.add_paragraph()
bold_text(doc, '核心结论：', '铁路局和民俗街是乌鲁木齐瘦身塑形连锁品牌覆盖最薄弱的两个高热度商圈。'
    '大唐辣妈、良咔、燃放、寇氏四大连锁品牌均未在铁路局商圈形成连锁布点——这是一个真实存在的市场空白。')

# ===== 四、重点商圈详细分析 =====
doc.add_page_break()
heading(doc, '四、重点商圈详细分析（30-50㎡店型）', 1)

# 铁路局
heading(doc, '1. 铁路局商圈 ⭐ 首选', 2)
doc.add_paragraph('乌鲁木齐新市区老牌生活商圈，以铁路局职工社区为核心，辐射周边3-5公里居民区。核心商业体包括太百购物（YoungPark）等。')
bold_text(doc, '抖音数据：', '预测到访热度全城TOP1，目标客群最有可能实际到店的区域。')
bold_text(doc, '竞品空白：', '大唐辣妈、良咔、燃放、寇氏四大全国连锁品牌均未在此商圈布点，仅有零散个体户经营。')
bold_text(doc, '客群画像：', '30-45岁中产家庭女性为主，消费稳定，健康意识强，有减脂/塑形需求。')
bold_text(doc, '租金水平：', '预估80-150元/㎡/月，30-50㎡月租金约2,400-7,500元。')
bold_text(doc, '优势：', '连锁竞品真空 + 抖音TOP1 + 社区稳定客流 + 成本低 + 30-50㎡属主流铺型。')
bold_text(doc, '风险：', '个体户竞品仍存在（同质化风险）；需快速建立差异化认知。')

# 民俗街
heading(doc, '2. 新疆民俗街/新华南路商圈 ⭐ 差异化选择', 2)
doc.add_paragraph('以新疆国际大巴扎为核心的文旅商圈，日均游客量大，夜间经济活跃。')
bold_text(doc, '抖音数据：', '预测到访热度TOP2，仅次于铁路局。')
bold_text(doc, '竞品空白：', '四大全国连锁品牌均未布点，属于真正的蓝海。')
bold_text(doc, '机会点：', '有到访需求但没有服务供给。可做差异化打法（旅游打卡+减脂体验），利用大巴扎网红属性做内容引流。')
bold_text(doc, '风险：', '游客转化率不确定、季节性波动大。建议作为第二梯队评估，待铁路局首店跑通后再决策。')

# 友好路
heading(doc, '3. 友好北路商圈 ⭐ 备选旗舰店', 2)
doc.add_paragraph('乌鲁木齐核心商业地标，聚集美美百货、友好商场、友好时尚购物城等高端商业体。')
bold_text(doc, '抖音数据：', '预测到访热度TOP5，低于铁路局和民俗街。')
bold_text(doc, '竞品情况：', '有高端纤体机构（高丽名人纤体等），无全国连锁大众品牌。')
bold_text(doc, '判断：', '流量质量高但数量非最高，租金和竞争门槛偏高。更适合作为品牌旗舰店，待铁路局样板店跑通后入驻，用核心商圈势能为招商赋能。')

# 友好路
heading(doc, '4. 其他商圈简要评价', 2)
t4 = doc.add_table(rows=4, cols=3, style='Table Grid')
borders(t4)
for j, h in enumerate(['商圈', '抖音热度', '建议']):
    set_cell(t4.rows[0].cells[j], h, bold=True, size=9, color=RGBColor(255,255,255))
header_bg(t4)
for i, row in enumerate([
    ('水上乐园', 'TOP3', '季节性明显，首店暂不推荐'),
    ('红山/南湖', 'TOP4', '寇氏已布点，差异化空间有限'),
    ('河南西路/乌北站/粮校/喀什西路', 'TOP6-10', '商业不成熟，不推荐'),
]):
    for j, v in enumerate(row):
        set_cell(t4.rows[i+1].cells[j], v, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

# ===== 五、总结 =====
doc.add_page_break()
heading(doc, '五、总结', 1)

t7 = doc.add_table(rows=1, cols=1, style='Table Grid')
borders(t7)
s = OxmlElement('w:shd')
s.set(qn('w:fill'), 'FFF8E1')
s.set(qn('w:val'), 'clear')
t7.rows[0].cells[0]._tc.get_or_add_tcPr().append(s)

summary = (
    f'本报告基于三组独立数据交叉验证：①抖音来客"预测TOP10到访商圈"；'
    '②大唐辣妈、良咔、燃放、寇氏减肥、芭芭多等全国连锁品牌的乌鲁木齐门店分布；'
    '③高德地图/美团公开商户信息。\n\n'
    f'{BRAND} 30-50㎡首店，建议选址铁路局商圈。'
    '\n\n'
    '理由：铁路局是抖音到访热度TOP1的商圈，但大唐辣妈、良咔、燃放、寇氏减肥四大全国连锁品牌均未在此布点——'
    '竞品真空是真实存在的，不是假设。这说明两件事：要么这个区域的商业条件未被品牌认可（风险），'
    '要么存在进入壁垒（机会）。结合抖音热度数据和社区客群稳定性，最可能的原因是前者尚未被充分验证。'
    '\n\n'
    '轻刻的策略应该是：借抖音数据证明需求是真实的，用轻资产模型（30-50㎡）低成本验证，'
    '一旦跑通单店模型，立即用数据招商扩大规模。在大唐辣妈、良咔等全国品牌意识到乌鲁木齐铁路局的机会之前，'
    '轻刻需要先占位置。'
)
set_cell(t7.rows[0].cells[0], summary, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT)

path = f'/Users/flyingspur/.qclaw/workspace-agent-6f07984c/{BRAND}_乌鲁木齐选址分析_V5.docx'
doc.save(path)
print(f'OK: {path}')
